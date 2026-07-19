"""Renders CVContent / CoverLetterContent into clean, ATS-friendly .docx files
using python-docx. Single column, standard fonts, no tables or text boxes —
this is deliberate: fancy layouts are the #1 cause of CVs getting mangled by
applicant tracking systems."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models import CoverLetterContent, CVContent

ACCENT = RGBColor(0x1B, 0x24, 0x30)  # ink navy — matches the app's own theme
BODY_FONT = "Calibri"


def _setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    # Bottom border as a simple divider (avoids using a table as a rule).
    pPr = p._p.get_or_add_pPr()
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B2430")
    pBdr.append(bottom)
    pPr.append(pBdr)


def render_cv_docx(cv: CVContent, out_path: Path) -> Path:
    doc = Document()
    _setup_page(doc)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name_p.add_run(cv.full_name)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    contact_bits = [b for b in [cv.location, cv.email, cv.phone, *cv.links] if b]
    if contact_bits:
        contact_p = doc.add_paragraph(" | ".join(contact_bits))
        contact_p.paragraph_format.space_after = Pt(6)
        for r in contact_p.runs:
            r.font.size = Pt(9.5)

    if cv.summary:
        _heading(doc, "Summary")
        doc.add_paragraph(cv.summary)

    if cv.skills:
        _heading(doc, "Skills")
        doc.add_paragraph(" • ".join(cv.skills))

    if cv.experience:
        _heading(doc, "Experience")
        for entry in cv.experience:
            line = doc.add_paragraph()
            line.paragraph_format.space_before = Pt(6)
            r1 = line.add_run(f"{entry.title}")
            r1.bold = True
            r2 = line.add_run(f"  —  {entry.company}")
            if entry.location:
                r2 = line.add_run(f" ({entry.location})")
            dates = f"{entry.start_date} – {entry.end_date}".strip(" –")
            if dates.strip():
                date_p = doc.add_paragraph(dates)
                for r in date_p.runs:
                    r.italic = True
                    r.font.size = Pt(9.5)
                date_p.paragraph_format.space_after = Pt(2)
            for bullet in entry.bullets:
                b = doc.add_paragraph(bullet, style="List Bullet")
                b.paragraph_format.space_after = Pt(1)

    if cv.education:
        _heading(doc, "Education")
        for edu in cv.education:
            line = f"{edu.degree}, {edu.institution}"
            if edu.year:
                line += f" ({edu.year})"
            p = doc.add_paragraph(line)
            if edu.details:
                d = doc.add_paragraph(edu.details, style="List Bullet")
                d.paragraph_format.space_after = Pt(1)

    if cv.certifications:
        _heading(doc, "Certifications")
        doc.add_paragraph(" • ".join(cv.certifications))

    if cv.projects:
        _heading(doc, "Projects")
        for proj in cv.projects:
            doc.add_paragraph(proj, style="List Bullet")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def render_cover_letter_docx(letter: CoverLetterContent, out_path: Path) -> Path:
    doc = Document()
    _setup_page(doc)

    name_p = doc.add_paragraph()
    run = name_p.add_run(letter.full_name)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = ACCENT

    if letter.contact_line:
        cp = doc.add_paragraph(letter.contact_line)
        for r in cp.runs:
            r.font.size = Pt(9.5)

    if letter.date:
        d = doc.add_paragraph(letter.date)
        d.paragraph_format.space_before = Pt(10)

    if letter.recipient or letter.company:
        rp = doc.add_paragraph()
        rp.paragraph_format.space_after = Pt(6)
        if letter.recipient:
            rp.add_run(letter.recipient + "\n")
        if letter.company:
            rp.add_run(letter.company)

    doc.add_paragraph(letter.salutation).paragraph_format.space_after = Pt(8)

    for para in letter.paragraphs:
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(8)

    doc.add_paragraph(letter.closing).paragraph_format.space_before = Pt(6)
    doc.add_paragraph(letter.full_name)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
